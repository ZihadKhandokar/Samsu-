from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches

from files.base_adapter import FileAdapter


class WordAdapter(FileAdapter):
    def create(self, path: Path, payload: Any) -> dict[str, Any]:
        if not isinstance(payload, dict):
            raise TypeError("Word creation requires a structured object.")
        document = Document()
        title = payload.get("title")
        if title:
            document.add_heading(str(title), level=0)
        for block in payload.get("blocks", []):
            self._add_block(document, block)
        document.save(path)
        return {
            "format": "docx",
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
        }

    @staticmethod
    def _add_block(document: Document, block: dict[str, Any]) -> None:
        kind = block.get("type", "paragraph")
        if kind == "heading":
            document.add_heading(str(block.get("text", "")), level=int(block.get("level", 1)))
        elif kind == "paragraph":
            paragraph = document.add_paragraph(str(block.get("text", "")))
            style = block.get("style")
            if style:
                paragraph.style = style
        elif kind == "table":
            rows = block.get("rows") or []
            if not rows:
                return
            width = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.style = block.get("style", "Table Grid")
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    table.cell(row_index, col_index).text = str(value)
        elif kind == "image":
            document.add_picture(
                str(block["path"]),
                width=Inches(float(block.get("width_inches", 5.5))),
            )
        elif kind == "page_break":
            document.add_page_break()
        else:
            raise ValueError(f"Unsupported Word block: {kind}")

    def read(self, path: Path) -> dict[str, Any]:
        document = Document(path)
        return {
            "format": "docx",
            "paragraphs": [p.text for p in document.paragraphs],
            "tables": [
                [[cell.text for cell in row.cells] for row in table.rows]
                for table in document.tables
            ],
        }

    def edit(self, path: Path, operations: list[dict[str, Any]]) -> dict[str, Any]:
        document = Document(path)
        for operation in operations:
            name = operation.get("operation")
            if name == "replace_text":
                old = str(operation.get("old_text", ""))
                new = str(operation.get("new_text", ""))
                if not old:
                    raise ValueError("replace_text requires old_text.")
                replacements = 0
                for paragraph in document.paragraphs:
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)
                        replacements += 1
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if old in cell.text:
                                cell.text = cell.text.replace(old, new)
                                replacements += 1
                if replacements == 0:
                    raise ValueError("Word text to replace was not found.")
            elif name == "append_paragraph":
                document.add_paragraph(str(operation.get("text", "")))
            elif name == "add_heading":
                document.add_heading(
                    str(operation.get("text", "")),
                    level=int(operation.get("level", 1)),
                )
            elif name == "add_page_break":
                document.add_page_break()
            elif name == "add_table":
                self._add_block(document, {"type": "table", **operation})
            else:
                raise ValueError(f"Unsupported Word operation: {name}")
        document.save(path)
        return {
            "format": "docx",
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
        }

    def verify(self, path: Path) -> dict[str, Any]:
        result = super().verify(path)
        document = Document(path)
        result.update(paragraphs=len(document.paragraphs), tables=len(document.tables))
        return result
